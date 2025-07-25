import os
import re
import unicodedata
import zipfile
from pathlib import Path
from io import BytesIO, StringIO
from datetime import datetime, date, timedelta

import streamlit as st
import pandas as pd
import docx
import pdfplumber
from openpyxl import load_workbook

# ==== DB Connection (Postgres vs SQLite) ====
IS_PG = "DATABASE_URL" in os.environ
if IS_PG:
    import psycopg2
    from urllib.parse import urlparse
    url = urlparse(os.environ["DATABASE_URL"])
    conn = psycopg2.connect(
        dbname=url.path[1:], user=url.username,
        password=url.password, host=url.hostname, port=url.port
    )
else:
    import sqlite3
    conn = sqlite3.connect("timesheets.db", check_same_thread=False)
c = conn.cursor()

# Ensure schema includes is_paid
c.execute("""
CREATE TABLE IF NOT EXISTS timesheet_entries (
    id SERIAL PRIMARY KEY,
    name TEXT,
    matched_as TEXT,
    ratio REAL,
    client TEXT,
    site_address TEXT,
    department TEXT,
    weekday_hours REAL,
    saturday_hours REAL,
    sunday_hours REAL,
    rate REAL,
    date_range TEXT,
    extracted_on TEXT,
    source_file TEXT,
    upload_timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    is_paid BOOLEAN DEFAULT FALSE
)
""")
conn.commit()

# ==== Helpers ====
def normalize_name(name: str) -> str:
    nfkd = unicodedata.normalize("NFKD", name)
    only_ascii = nfkd.encode("ASCII", "ignore").decode("utf-8")
    return re.sub(r"[^a-zA-Z]", "", only_ascii).lower()

def load_rate_database(source):
    custom, normed, to_raw = {}, {}, {}
    wb = load_workbook(source, data_only=True)
    for sheet in wb.sheetnames:
        df0 = pd.read_excel(source, sheet_name=sheet, header=None)
        header_row = next(
            (i for i,v in enumerate(df0.iloc[:,0])
             if isinstance(v,str)
                and v.strip().lower()=="name"
                and isinstance(df0.iat[i,1],str)
                and df0.iat[i,1].strip().lower()=="pay rate"),
            None
        )
        if header_row is None: continue
        df = pd.read_excel(source, sheet_name=sheet, header=header_row)
        if "Name" not in df.columns or "Pay Rate" not in df.columns: continue
        df = df[["Name","Pay Rate"]].dropna()
        df["Pay Rate"] = pd.to_numeric(df["Pay Rate"], errors="coerce")
        for _,r in df.iterrows():
            raw = str(r["Name"]).strip()
            rate = float(r["Pay Rate"])
            custom[raw] = rate
            n = normalize_name(raw)
            normed[n] = rate
            to_raw[n] = raw
    return custom, normed, to_raw

def lookup_match(name: str):
    n = normalize_name(name)
    if n in normalized_rates:
        return norm_to_raw[n], normalized_rates[n], 1.0
    return name, 15.0, 0.0

# ==== Sidebar: Rate‑Sheet Upload ====
RATE_FILE = "pay_rates.xlsx"
rate_uploads = st.sidebar.file_uploader(
    "➕ Upload pay‑rate XLSX(s)",
    type=["xlsx"], accept_multiple_files=True
)
custom_rates, normalized_rates, norm_to_raw = {}, {}, {}
def _merge(src):
    cr,nr,nt = load_rate_database(src)
    custom_rates.update(cr)
    normalized_rates.update(nr)
    norm_to_raw.update(nt)

if rate_uploads:
    for f in rate_uploads: _merge(f)
    st.sidebar.success(f"Merged {len(rate_uploads)} rate sheet(s).")
elif Path(RATE_FILE).exists():
    try:
        _merge(RATE_FILE)
        st.sidebar.info(f"Loaded local `{RATE_FILE}`.")
    except Exception as e:
        st.sidebar.error(f"Error loading `{RATE_FILE}`: {e}")
if not normalized_rates:
    st.sidebar.warning("No pay‑rates found; defaulting to £15/hr.")

# ==== Extraction Logic ====
def extract_from_docx(file) -> list[dict]:
    doc = docx.Document(file)
    if not doc.tables: return []
    tbl = doc.tables[0]
    hdr = tbl.rows[0].cells[0].text.split("\n")
    hdr = [h.strip() for h in hdr if h.strip()]
    client = name = site = None
    for i,line in enumerate(hdr):
        low=line.lower()
        if low.startswith("client"):
            parts=line.split(None,1)
            client = parts[1].strip() if len(parts)>1 else ""
            if i+1<len(hdr): name=hdr[i+1].strip()
        if low.startswith("site address"):
            parts=line.split("\t",1)
            site = parts[1].strip() if len(parts)>1 else ""
    header_row = next(
        (i for i,row in enumerate(tbl.rows)
         if row.cells[0].text.strip().lower()=="date"), 
        None
    )
    if header_row is None: return []
    date_re = re.compile(r"\d{2}\.\d{2}\.\d{4}")
    wd=sa=su=0.0; dates=[]
    for row in tbl.rows[header_row+1:]:
        dt=row.cells[0].text.strip()
        if not date_re.match(dt): continue
        dates.append(dt)
        day=row.cells[1].text.strip().lower()
        try: hrs=float(row.cells[4].text.strip())
        except: hrs=0.0
        if day=="saturday": sa+=hrs
        elif day=="sunday": su+=hrs
        else: wd+=hrs
    if not dates: return []
    dr = f"{min(dates)}–{max(dates)}"
    matched,rate,ratio = lookup_match(name or "")
    return [{
        "id": None,
        "name": name or "",
        "matched_as": matched,
        "ratio": ratio,
        "client": client or "",
        "site_address": site or "",
        "department": "",
        "weekday_hours": wd,
        "saturday_hours": sa,
        "sunday_hours": su,
        "rate": rate,
        "date_range": dr,
        "extracted_on": datetime.now().isoformat(),
        "source_file": None,
        "is_paid": False
    }]

def extract_from_pdf(file) -> list[dict]:
    return []

# ==== Sidebar: Timesheet Upload ====
st.sidebar.header("Upload Timesheets")
st.sidebar.markdown("""
1. Upload **.docx**, **.pdf** or **.zip**  
2. Check duplicates & mark paid  
3. Download Excel  
""")
uploaded = st.sidebar.file_uploader(
    "Choose timesheet file(s)", accept_multiple_files=True
)

# ==== Main Tabs ====
tabs = st.tabs([
    "Upload & Review",
    "History",
    "Matches",
    "Dashboard",
    "Settings",
    "BrightPay"
])

# ---- 1) Upload & Review ----
with tabs[0]:
    st.header("📤 Upload & Review Timesheets")
    if not uploaded:
        st.info("Upload .docx/.pdf/.zip to begin.")
    else:
        progress = st.progress(0)
        summaries = []
        for i,uf in enumerate(uploaded):
            st.write(f"➡️ {uf.name} ({i+1}/{len(uploaded)})")
            lower = uf.name.lower()
            def handle(recs):
                for r in recs:
                    r["source_file"] = uf.name
                    summaries.append(r)
            if lower.endswith(".zip"):
                try:
                    z = zipfile.ZipFile(uf)
                    for m in [m for m in z.namelist() if m.lower().endswith((".docx",".pdf"))]:
                        buf = BytesIO(z.read(m)); buf.name=m
                        recs = (extract_from_docx(buf) if m.lower().endswith(".docx")
                                else extract_from_pdf(buf))
                        st.write(f" • {m}: {len(recs)} rec(s)")
                        handle(recs)
                except zipfile.BadZipFile:
                    st.error(f"{uf.name} invalid ZIP.")
            elif lower.endswith(".docx"):
                recs = extract_from_docx(uf); st.write(f" • DOCX: {len(recs)} rec(s)"); handle(recs)
            elif lower.endswith(".pdf"):
                recs = extract_from_pdf(uf); st.write(f" • PDF: {len(recs)} rec(s)"); handle(recs)
            else:
                st.warning(f"Unsupported: {uf.name}")
            progress.progress((i+1)/len(uploaded))

        if not summaries:
            st.error("No records extracted.")
        else:
            df = pd.DataFrame(summaries)
            with st.expander("🔍 Raw summaries"):
                st.dataframe(df.drop(columns=["id"]), use_container_width=True)

            # Duplicate check
            existing,new = [],[]
            for r in summaries:
                if IS_PG:
                    c.execute("SELECT COUNT(*) FROM timesheet_entries WHERE name=%s AND date_range=%s",
                              (r["name"],r["date_range"]))
                else:
                    c.execute("SELECT COUNT(*) FROM timesheet_entries WHERE name=? AND date_range=?",
                              (r["name"],r["date_range"]))
                cnt = c.fetchone()[0]
                (existing if cnt>0 else new).append(r)

            if existing:
                st.warning("⚠️ Duplicates skipped:")
                st.dataframe(pd.DataFrame(existing)[["name","date_range","source_file"]], use_container_width=True)

            # Excel export
            buf = BytesIO()
            with pd.ExcelWriter(buf, engine="openpyxl") as w:
                df.drop(columns=["id"]).to_excel(w, index=False)
            buf.seek(0)
            st.download_button("📥 Download All Summaries", data=buf,
                file_name=f"summaries_{date.today()}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

            # Persist new
            if new:
                ph = ",".join("%s" if IS_PG else "?" for _ in range(13))
                sql = f"""
                    INSERT INTO timesheet_entries
                      (name,matched_as,ratio,client,site_address,department,
                       weekday_hours,saturday_hours,sunday_hours,rate,
                       date_range,extracted_on,source_file)
                    VALUES({ph})
                """
                for r in new:
                    params = [
                        r["name"],r["matched_as"],r["ratio"],
                        r["client"],r["site_address"],r["department"],
                        r["weekday_hours"],r["saturday_hours"],
                        r["sunday_hours"],r["rate"],
                        r["date_range"],r["extracted_on"],r["source_file"]
                    ]
                    c.execute(sql, params)
                conn.commit()
                st.success(f"Inserted {len(new)} new rec(s).")
            else:
                st.info("No new records to insert.")

# ---- 2) History ----
with tabs[1]:
    st.header("🗃️ Upload History")
    st.markdown("Presets or custom range:")
    presets = ["Last 30 days","This month","Last month","Year to date","Custom"]
    choice = st.selectbox("Range", presets)
    today = date.today()
    if choice != "Custom":
        if choice=="Last 30 days":
            start,end = today-timedelta(30),today
        elif choice=="This month":
            start,end = today.replace(day=1),today
        elif choice=="Last month":
            f=today.replace(day=1); l=f-timedelta(1)
            start,end = l.replace(day=1),l
        else:
            start,end = date(today.year,1,1),today
        st.write(f"**{choice}**: {start} → {end}")
    else:
        start,end = st.date_input("Custom range",
                                  value=(today-timedelta(30),today),
                                  min_value=date(2020,1,1),max_value=today)

    c.execute("""
        SELECT id,name,matched_as,ratio,client,site_address,department,
               weekday_hours,saturday_hours,sunday_hours,rate AS rate,
               date_range,extracted_on,source_file,upload_timestamp,is_paid
        FROM timesheet_entries
        ORDER BY upload_timestamp DESC
    """)
    hist = pd.DataFrame(c.fetchall(), columns=[
        "id","Name","Matched As","Ratio","Client","Site Address",
        "Department","Weekday Hours","Saturday Hours","Sunday Hours",
        "Rate (£)","Date Range","Extracted On","Source File",
        "Upload Timestamp","Paid?"
    ])
    hist["Upload Timestamp"] = pd.to_datetime(hist["Upload Timestamp"]).dt.date
    view = hist[(hist["Upload Timestamp"]>=start)&(hist["Upload Timestamp"]<=end)]
    if view.empty:
        st.info("No entries in this range.")
    else:
        # only include persisted rows with real IDs
        valid = view.dropna(subset=["id"]).copy()
        labels = valid.apply(
            lambda r: f"{int(r['id'])}: {r['Name']} ({r['Date Range']}) Paid? {r['Paid?']}",
            axis=1
        )
        # convert to plain Python list
        label_list = labels.values.tolist()

        selected = st.multiselect("Select entries", label_list)
        sel_ids = [int(s.split(":")[0]) for s in selected]

        c1,c2,c3 = st.columns(3)
        if sel_ids and c1.button("Delete selected"):
            ph=",".join("%s" if IS_PG else "?" for _ in sel_ids)
            c.execute(f"DELETE FROM timesheet_entries WHERE id IN ({ph})", sel_ids)
            conn.commit()
            st.success(f"Deleted {len(sel_ids)} record(s).")
        if sel_ids and c2.button("Export selected"):
            df_sel = valid[valid["id"].isin(sel_ids)]
            buf2 = BytesIO()
            with pd.ExcelWriter(buf2, engine="openpyxl") as w:
                df_sel.to_excel(w, index=False)
            buf2.seek(0)
            st.download_button("Download Export", data=buf2,
                file_name=f"export_{date.today()}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        if sel_ids and c3.button("Mark paid"):
            ph=",".join("%s" if IS_PG else "?" for _ in sel_ids)
            c.execute(f"UPDATE timesheet_entries SET is_paid=TRUE WHERE id IN ({ph})", sel_ids)
            conn.commit()
            st.success(f"Marked {len(sel_ids)} paid.")

# ---- 3) Matches ----
with tabs[2]:
    st.header("🔗 Name → Rate Matches")
    st.markdown("Inline edits:")
    c.execute("SELECT DISTINCT name,matched_as,ratio FROM timesheet_entries ORDER BY name")
    dfm = pd.DataFrame(c.fetchall(), columns=["Timesheet Name","Matched Rate Name","Confidence"])
    if dfm.empty:
        st.info("No matches yet.")
    else:
        edited = st.data_editor(
            dfm,
            column_config={
                "Timesheet Name":{"disabled":True},
                "Matched Rate Name":{},
                "Confidence":{"min_value":0.0,"max_value":1.0}
            },
            use_container_width=True
        )
        if st.button("Save match edits"):
            diffs = edited.merge(dfm,indicator=True,how="outer").query("_merge!='both'")
            for _,r in diffs.iterrows():
                nm,cf = r["Matched Rate Name"],r["Confidence"]
                name = r["Timesheet Name"]
                if IS_PG:
                    c.execute("UPDATE timesheet_entries SET matched_as=%s,ratio=%s WHERE name=%s",(nm,cf,name))
                else:
                    c.execute("UPDATE timesheet_entries SET matched_as=?,ratio=? WHERE name=?",(nm,cf,name))
            conn.commit()
            st.success(f"Updated {len(diffs)} match(es).")

# ---- 4) Dashboard ----
with tabs[3]:
    st.header("📊 Dashboard")
    st.markdown("… your charts here …")

# ---- 5) Settings ----
with tabs[4]:
    st.header("⚙️ Settings & Info")
    st.markdown("""
    - History multiselect now uses `labels.values.tolist()`.  
    - No more `None` selections or `tolist()` errors.  
    """)

# ---- 6) BrightPay Export ----
with tabs[5]:
    st.header("💼 Export to BrightPay")
    st.markdown("1) Download mapping template  2) Fill Employee IDs  3) Re‑upload")

    # Template download
    c.execute("SELECT DISTINCT name FROM timesheet_entries ORDER BY name")
    names = [r[0] for r in c.fetchall()]
    tmpl = pd.DataFrame({"Name": names, "Employee ID": [""]*len(names)})
    tmpl_buf = BytesIO()
    with pd.ExcelWriter(tmpl_buf, engine="openpyxl") as w:
        tmpl.to_excel(w, index=False, sheet_name="Mapping")
    tmpl_buf.seek(0)
    st.download_button("📄 Download Mapping Template", data=tmpl_buf,
        file_name="brightpay_mapping_template.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    # Upload filled mapping
    emp_map = st.file_uploader("🔄 Upload filled mapping", type=["xlsx","csv"])
    bp_df = None
    if emp_map:
        emp = pd.read_excel(emp_map) if emp_map.name.lower().endswith("xlsx") else pd.read_csv(emp_map)
        if {"Name","Employee ID"}.issubset(emp.columns):
            c.execute("""
                SELECT name,weekday_hours,saturday_hours,sunday_hours,rate,date_range,client
                FROM timesheet_entries ORDER BY upload_timestamp DESC
            """)
            bp_df = pd.DataFrame(c.fetchall(), columns=[
                "Name","WD","Sat","Sun","Rate","Period","Client"
            ]).merge(emp[["Name","Employee ID"]], on="Name", how="left")
            missing = bp_df[bp_df["Employee ID"].isna()]["Name"].unique()
            if len(missing):
                st.warning(f"No Emp ID for: {missing.tolist()}")
        else:
            st.error("Mapping must have 'Name' & 'Employee ID' cols.")

    pay_el = st.selectbox("Pay Element", ["Standard Hours","Overtime","Holiday"])
    if bp_df is not None and st.button("📥 Generate BrightPay CSV"):
        out=[]
        for _,row in bp_df.iterrows():
            start,end = row["Period"].split("–")
            if row["WD"]>0:
                out.append({
                    "Employee ID":row["Employee ID"],
                    "Period Start":start,
                    "Period End":end,
                    "Pay Element":pay_el,
                    "Units":row["WD"],
                    "Rate":row["Rate"],
                    "Cost Center":row["Client"]
                })
        if out:
            csv_buf = StringIO()
            pd.DataFrame(out).to_csv(csv_buf,index=False)
            st.download_button("📂 Download BrightPay CSV", data=csv_buf.getvalue(),
                file_name=f"brightpay_{date.today().isoformat()}.csv", mime="text/csv")
        else:
            st.info("No hours to export.")
