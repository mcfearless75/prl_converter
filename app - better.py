import streamlit as st
import pandas as pd
import docx
import pdfplumber
import re
import unicodedata
from io import BytesIO
from datetime import datetime
from pathlib import Path
from openpyxl import Workbook, load_workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Font, PatternFill
from difflib import SequenceMatcher

# ──────────────────────────────────────────────────────────────────────────────
# 0) PAGE CONFIGURATION
# ──────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="PRL Timesheet Portal",
    page_icon="📑",
    layout="wide"
)

# ──────────────────────────────────────────────────────────────────────────────
# 1) LOAD “pay details.xlsx” FROM ALL SHEETS, BUILD NAME→RATE MAPPINGS
# ──────────────────────────────────────────────────────────────────────────────

RATE_FILE_PATH = "pay details.xlsx"
DEFAULT_RATE = 15.0
SIMILARITY_THRESHOLD = 0.60  # Adjust if needed for fuzzy matches

def normalize_name(s: str) -> str:
    """ Lowercase, strip accents, remove punctuation, collapse spaces. """
    s = s.lower().strip()
    s = unicodedata.normalize("NFD", s)
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
    s = re.sub(r"[^a-z0-9\s]", "", s)
    s = re.sub(r"\s+", " ", s)
    return s

@st.cache_data
def load_rate_database(excel_path: str):
    """
    1) Iterate all sheets in pay details.xlsx.
    2) In each sheet, find the header row where A="Name" and B="Pay Rate".
    3) Read that sheet with header=header_row.
    4) Keep only rows with valid “Name” & numeric “Pay Rate”.
    5) Build:
       • custom_rates: {raw_name → day_rate}
       • normalized_rates: {normalize_name(raw_name) → day_rate}
       • norm_to_raw:      {normalize_name(raw_name) → raw_name}
    6) Return (custom_rates, normalized_rates, normalized_keys, norm_to_raw).
    """
    wb = load_workbook(excel_path, read_only=True)
    custom_rates = {}
    normalized_rates = {}
    norm_to_raw = {}

    for sheet in wb.sheetnames:
        df_raw = pd.read_excel(excel_path, sheet_name=sheet, header=None)
        header_row = None
        # Find the header row where column A is "Name" and B is "Pay Rate"
        for idx, val in enumerate(df_raw[0]):
            if isinstance(val, str) and val.strip().lower() == "name":
                second = df_raw.iat[idx, 1]
                if isinstance(second, str) and second.strip().lower() == "pay rate":
                    header_row = idx
                    break
        if header_row is None:
            continue

        df = pd.read_excel(excel_path, sheet_name=sheet, header=header_row)
        if "Name" not in df.columns or "Pay Rate" not in df.columns:
            continue

        df = df[["Name", "Pay Rate"]].copy()
        df["Pay Rate"] = pd.to_numeric(df["Pay Rate"], errors="coerce")
        df = df.dropna(subset=["Name", "Pay Rate"])

        for _, row in df.iterrows():
            raw_name = str(row["Name"]).strip()
            rate = float(row["Pay Rate"])
            custom_rates[raw_name] = rate
            norm = normalize_name(raw_name)
            normalized_rates[norm] = rate
            norm_to_raw[norm] = raw_name

    normalized_keys = list(normalized_rates.keys())
    return custom_rates, normalized_rates, normalized_keys, norm_to_raw

# Sidebar: reload rates button
if st.sidebar.button("🔄 Reload Pay Rates"):
    st.cache_data.clear()
    st.experimental_rerun()

custom_rates, normalized_rates, normalized_keys, norm_to_raw = load_rate_database(RATE_FILE_PATH)

# ──────────────────────────────────────────────────────────────────────────────
# 2) FUZZY LOOKUP AND PAY CALCULATION
# ──────────────────────────────────────────────────────────────────────────────
def lookup_match(name: str):
    """
    1) Normalize the extracted name.
    2) If exact normalized key exists → return (raw_name, rate, 1.0).
    3) Otherwise, iterate all normalized_keys, compute SequenceMatcher ratio.
       - If best_ratio ≥ SIMILARITY_THRESHOLD, use that match.
    4) If still no good match, return (None, DEFAULT_RATE, best_ratio).
    """
    norm = normalize_name(name)
    if not norm:
        return None, DEFAULT_RATE, 0.0

    if norm in normalized_rates:
        raw = norm_to_raw[norm]
        return raw, normalized_rates[norm], 1.0

    best_ratio = 0.0
    best_norm = None
    for nk in normalized_keys:
        ratio = SequenceMatcher(None, nk, norm).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_norm = nk

    if best_norm and best_ratio >= SIMILARITY_THRESHOLD:
        raw = norm_to_raw[best_norm]
        return raw, normalized_rates[best_norm], best_ratio

    return None, DEFAULT_RATE, best_ratio

def hhmm_to_float(hhmm: str) -> float:
    try:
        h, m = hhmm.strip().split(":")
        return int(h) + int(m) / 60.0
    except:
        return 0.0

def calculate_pay(name: str, daily_data: list[dict]):
    """
    Overtime & weekend rules:
      - Mon–Fri summed into weekday_hours.
      - First 50 weekday_hours at rate, any beyond 50 at 1.5×.
      - Saturday hours always 1.5×.
      - Sunday hours always 1.75×.
    Returns (weekday_hours, sat_hours, sun_hours, rate, total_pay, matched_raw_name, ratio).
    """
    matched_raw, rate, ratio = lookup_match(name)

    weekday_hours = 0.0
    sat_hours = 0.0
    sun_hours = 0.0

    for e in daily_data:
        h = e["hours"]
        wd = e["weekday"]
        if wd == "Saturday":
            sat_hours += h
        elif wd == "Sunday":
            sun_hours += h
        else:
            weekday_hours += h

    overtime = max(0.0, weekday_hours - 50.0)
    regular_wd = weekday_hours - overtime

    pay_regular = regular_wd * rate
    pay_overtime = overtime * rate * 1.5
    pay_sat = sat_hours * rate * 1.5
    pay_sun = sun_hours * rate * 1.75

    total_pay = pay_regular + pay_overtime + pay_sat + pay_sun

    return weekday_hours, sat_hours, sun_hours, rate, total_pay, matched_raw, ratio

# ──────────────────────────────────────────────────────────────────────────────
# 3) DOCX EXTRACTION
# ──────────────────────────────────────────────────────────────────────────────
def extract_timesheet_data(file) -> dict:
    """
    Parse a .docx:
     1) Scan tables for any cell containing “Client”; extract the next ALL‐CAPS line as Name.
     2) In the same table(s), collect daily hours (col index 4) keyed by weekday in col index 1.
     3) Also scan cells for “Site Address: …”.
     4) Fallback 1: reversed paragraphs → first ALL‐CAPS line as Name.
     5) Fallback 2: filename if still empty.
     6) Call calculate_pay(...) to get (weekday_hours, sat_hours, sun_hours, rate, pay, matched, ratio).
     7) Return a dict with all fields + “Source File”.
    """
    doc = docx.Document(file)
    name = client = site_address = ""
    date_list = []
    daily_data = []

    # 3.1) Find “Client” cell → extract client & name beneath
    for table in doc.tables:
        found_c = False
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text or ""
                if "Client" in txt:
                    found_c = True
                    lines = [ln.strip() for ln in txt.splitlines() if ln.strip()]
                    for idx, ln in enumerate(lines):
                        if ln.lower().startswith("client"):
                            parts = re.split(r"Client[:\-\s]+", ln, flags=re.IGNORECASE)
                            if len(parts) > 1:
                                client = parts[1].strip()
                            if idx + 1 < len(lines):
                                cand = lines[idx + 1]
                                if cand == cand.upper() and len(cand.split()) >= 2:
                                    name = cand.title()
                            break
                    break
            if found_c:
                break
        if found_c:
            break

    # 3.2) Collect hours, site address, date list
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 5:
                hrs_txt = cells[4].text.strip()
                day_txt = cells[1].text.strip()
                date_txt = cells[0].text.strip()
                if hrs_txt and hrs_txt not in ["-", "–", "—"] and day_txt:
                    try:
                        val = float(hrs_txt)
                        daily_data.append({"weekday": day_txt, "hours": val})
                    except:
                        pass
                if re.match(r"\d{2}\.\d{2}\.\d{4}", date_txt):
                    try:
                        d_obj = datetime.strptime(date_txt, "%d.%m.%Y")
                        date_list.append(d_obj)
                    except:
                        pass
            for cell in cells:
                txt = (cell.text or "").strip()
                if "Site Address" in txt and not site_address:
                    m = re.search(r"Site Address[:\-\s]*(.+)", txt)
                    if m:
                        site_address = m.group(1).strip()

    # 3.3) Fallback 1: reversed paragraphs for ALL‐CAPS name
    if not name:
        for para in reversed(doc.paragraphs):
            text = (para.text or "").strip()
            if text and text == text.upper() and len(text.split()) >= 2 and "PRL" not in text:
                name = text.title()
                break

    # 3.4) Fallback 2: filename as name
    if not name:
        stem = Path(file.name).stem
        name = stem.replace("_", " ").replace("-", " ").title()

    wd_hrs, sat_hrs, sun_hrs, rate, total_pay, matched_raw, ratio = calculate_pay(name, daily_data)

    date_range = ""
    if date_list:
        ds = sorted(date_list)
        date_range = f"{ds[0].strftime('%d.%m.%Y')}–{ds[-1].strftime('%d.%m.%Y')}"

    return {
        "Name": name,
        "Matched As": matched_raw or "No match",
        "Ratio": round(ratio, 2),
        "Client": client,
        "Site Address": site_address,
        "Department": "",
        "Weekday Hours": wd_hrs,
        "Saturday Hours": sat_hrs,
        "Sunday Hours": sun_hrs,
        "Rate (£)": rate,
        "Date Range": date_range,
        "Extracted On": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "Source File": file.name
    }

# ──────────────────────────────────────────────────────────────────────────────
# 4) PDF EXTRACTION
# ──────────────────────────────────────────────────────────────────────────────
def extract_timesheet_data_pdf(file) -> list[dict]:
    """
    Parse an Anel PDF:
     1) Extract “Report Range: … to …” from page1 for date_range.
     2) Find header “ID Name Paylink” + “Tot Sat-Sun”.
     3) For each line until “Grand Totals”:
        - tokens = line.split()
        - last 9 tokens are H:MM → daily_data
        - dash_idx = tokens.index("-")
        - Name = tokens[1:dash_idx]
        - Build daily_data from Mon–Sun.
        - Call calculate_pay(name, daily_data).
        - Append a dict with fields including “Weekday Hours”, “Saturday Hours”, “Sunday Hours”.
    4) Return list of dicts.
    """
    with pdfplumber.open(file) as pdf:
        page1 = pdf.pages[0]
        raw = page1.extract_text()
    lines = raw.split("\n")

    date_range = ""
    for line in lines:
        if line.startswith("Report Range:"):
            m = re.search(
                r"Report Range:\s*(?:Mon|Tue|Wed|Thu|Fri|Sat|Sun)\s+(\d{2}/\d{2}/\d{2})\s+to\s+(?:Mon|Tue|Wed|Thu|Fri|Sat|Sun)\s+(\d{2}/\d{2}/\d{2})",
                line
            )
            if m:
                def fmt(d: str) -> str:
                    dt = datetime.strptime(d, "%d/%m/%y")
                    return dt.strftime("%d.%m.%Y")
                date_range = f"{fmt(m.group(1))}–{fmt(m.group(2))}"
            break

    header_idx = None
    for idx, line in enumerate(lines):
        if "ID Name Paylink" in line and "Tot Sat-Sun" in line:
            header_idx = idx
            break

    if header_idx is None:
        return [{
            "Name": "",
            "Matched As": "No match",
            "Ratio": 0.0,
            "Client": "",
            "Site Address": "",
            "Department": "",
            "Weekday Hours": 0.0,
            "Saturday Hours": 0.0,
            "Sunday Hours": 0.0,
            "Rate (£)": 0.0,
            "Date Range": date_range,
            "Extracted On": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "Source File": file.name
        }]

    results: list[dict] = []
    time_re = re.compile(r"^\d{1,2}:\d{2}$")

    for raw_line in lines[header_idx+1:]:
        if raw_line.strip().startswith("Grand Totals"):
            break
        line = raw_line.strip()
        if not line:
            continue

        tokens = line.split()
        n = len(tokens)
        if n < 10:
            continue

        block = tokens[-9:]
        if not all(time_re.match(t) for t in block):
            continue

        try:
            dash_idx = tokens.index("-")
        except ValueError:
            continue

        name = " ".join(tokens[1:dash_idx]).title()
        daily_data = []
        for wd, idx_tok in [
            ("Monday", 0), ("Tuesday", 1), ("Wednesday", 2),
            ("Thursday", 3), ("Friday", 4), ("Saturday", 6), ("Sunday", 7)
        ]:
            h = hhmm_to_float(block[idx_tok])
            if h > 0:
                daily_data.append({"weekday": wd, "hours": h})

        wd_hrs, sat_hrs, sun_hrs, rate, total_pay, matched_raw, ratio = calculate_pay(name, daily_data)

        comp = ""
        if dash_idx + 2 < n:
            comp = " ".join(tokens[dash_idx+1 : dash_idx+3])

        site_addr = ""
        if dash_idx + 5 < n:
            site_addr = " ".join(tokens[dash_idx+3 : dash_idx+5])

        dept = ""
        dept_start = dash_idx + 5
        dept_end = n - 9
        if dept_start < dept_end:
            dept = " ".join(tokens[dept_start : dept_end]).rstrip("-")

        results.append({
            "Name": name,
            "Matched As": matched_raw or "No match",
            "Ratio": round(ratio, 2),
            "Client": comp,
            "Site Address": site_addr,
            "Department": dept,
            "Weekday Hours": wd_hrs,
            "Saturday Hours": sat_hrs,
            "Sunday Hours": sun_hrs,
            "Rate (£)": rate,
            "Date Range": date_range,
            "Extracted On": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "Source File": file.name
        })

    if not results:
        return [{
            "Name": "",
            "Matched As": "No match",
            "Ratio": 0.0,
            "Client": "",
            "Site Address": "",
            "Department": "",
            "Weekday Hours": 0.0,
            "Saturday Hours": 0.0,
            "Sunday Hours": 0.0,
            "Rate (£)": 0.0,
            "Date Range": date_range,
            "Extracted On": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "Source File": file.name
        }]

    return results

# ──────────────────────────────────────────────────────────────────────────────
# 5) LAYOUT: SIDEBAR + MAIN AREA WITH READ-ONLY TABLE + EXPORT
# ──────────────────────────────────────────────────────────────────────────────

# Sidebar header & instructions
st.sidebar.title("PRL Timesheet Portal")
st.sidebar.markdown(
    """
    1. Upload your **Word (.docx)** or **Anel PDF (.pdf)** timesheet files.  
    2. Confirm extracted names vs. pay-detail matches (expand Debug below).  
    3. If any names are unmatched, use the **Manual Rate Override** section.  
    4. Once correct, review the table and export with formulas.
    """
)

# File uploader in sidebar
uploaded_files = st.sidebar.file_uploader(
    "📂 Upload Timesheets",
    type=["docx", "pdf"],
    accept_multiple_files=True
)

# Manual override expander in sidebar (initially collapsed)
override_exp = st.sidebar.expander("✏️ Manual Rate Override", expanded=False)

# Only show the main UI once files are uploaded
if uploaded_files:
    # 5.1) Extract & consolidate all rows
    all_rows = []
    for file in uploaded_files:
        lower = file.name.lower()
        if lower.endswith(".docx"):
            rec = extract_timesheet_data(file)
            if not rec["Name"]:
                stem = Path(file.name).stem
                rec["Name"] = stem.replace("_", " ").replace("-", " ").title()
            all_rows.append(rec)

        elif lower.endswith(".pdf"):
            pdf_list = extract_timesheet_data_pdf(file)
            for r in pdf_list:
                if not r["Name"]:
                    continue
                all_rows.append(r)

    df = pd.DataFrame(all_rows)

    # 5.2) Debug: show extracted vs. matched table in an expander
    st.markdown("## 🔎 Extracted vs. Matched Pay-Detail Entries")
    debug_df = (
        df[["Name", "Matched As", "Ratio", "Rate (£)", "Source File"]]
        .drop_duplicates()
        .reset_index(drop=True)
    )
    with st.expander("Show Name-Match Debug Table", expanded=False):
        st.dataframe(debug_df, use_container_width=True)

    # Show names that had no match (Ratio < threshold)
    no_match = debug_df.loc[debug_df["Matched As"] == "No match", "Name"].tolist()
    if no_match:
        st.error(
            "⚠️ The following names did not match any pay-detail entry "
            f"(ratio < {SIMILARITY_THRESHOLD:.2f}):\n\n"
            + "\n".join(f"- {nm}" for nm in no_match)
        )

    # 5.3) Manual override UI in sidebar
    unmatched = debug_df.loc[debug_df["Matched As"] == "No match", "Name"].tolist()
    if unmatched:
        with override_exp:
            st.markdown("**Map an extracted name to the correct pay-detail entry:**")
            override_name = st.selectbox("Unmatched Name", options=[""] + unmatched)
            if override_name:
                chosen = st.selectbox(
                    "Select Pay-Detail Contractor",
                    options=[""] + list(custom_rates.keys())
                )
                if chosen and st.button("Apply Override"):
                    correct_rate = custom_rates[chosen]
                    df.loc[df["Name"] == override_name, "Matched As"] = chosen
                    df.loc[df["Name"] == override_name, "Rate (£)"] = correct_rate
                    # Recompute pay breakdown for those rows
                    def recalc_total_pay(row):
                        wd = row["Weekday Hours"]
                        sat = row["Saturday Hours"]
                        sun = row["Sunday Hours"]
                        rate_override = correct_rate
                        overtime = max(0.0, wd - 50.0)
                        regular_wd = wd - overtime
                        pay_regular = regular_wd * rate_override
                        pay_overtime = overtime * rate_override * 1.5
                        pay_sat = sat * rate_override * 1.5
                        pay_sun = sun * rate_override * 1.75
                        return pay_regular + pay_overtime + pay_sat + pay_sun
                    df.loc[df["Name"] == override_name, "Calculated Pay (£)"] = df[df["Name"] == override_name].apply(recalc_total_pay, axis=1)
                    st.success(f"✅ Overrode rate for '{override_name}' → £{correct_rate:.2f}")

    # 5.4) Read‐only DataFrame in main area
    st.markdown("## 📋 Final Timesheet Table (Read‐Only)")
    # Compute initial "Calculated Pay (£)" if missing
    if "Calculated Pay (£)" not in df.columns:
        df["Calculated Pay (£)"] = df.apply(
            lambda row: calculate_pay(
                row["Name"],
                []  # daily_data isn't needed because we store breakdown columns
            )[4],
            axis=1
        )
    st.dataframe(df, use_container_width=True)

    # 5.5) Weekly summary based on df
    summary_df = (
        df.groupby("Matched As")[["Calculated Pay (£)", "Weekday Hours", "Saturday Hours", "Sunday Hours"]]
        .sum()
        .reset_index()
        .rename(columns={"Matched As": "Name"})
    )

    st.markdown("## 💰 Weekly Summary")
    col1, col2 = st.columns([3, 1])
    with col1:
        st.dataframe(summary_df, use_container_width=True)
    with col2:
        total_hours_all = (
            summary_df["Weekday Hours"].sum()
            + summary_df["Saturday Hours"].sum()
            + summary_df["Sunday Hours"].sum()
        )
        total_pay_all = summary_df["Calculated Pay (£)"].sum()
        st.metric(label="Total Hours This Period", value=f"{total_hours_all:.2f}")
        st.metric(label="Total Pay This Period", value=f"£{total_pay_all:.2f}")

    # 5.6) Export section: Excel with formulas
    st.markdown("---")
    st.markdown("### 📥 Download Final Report (Excel with Formulas)")
    output = BytesIO()
    wb_out = Workbook()
    ws = wb_out.active
    ws.title = "Timesheets"

    # Build header row with additional columns for formulas
    headers = [
        "Name", "Matched As", "Ratio", "Client", "Site Address", "Department",
        "Weekday Hours", "Saturday Hours", "Sunday Hours", "Rate (£)",
        "Regular Pay (£)", "Overtime Pay (£)", "Saturday Pay (£)", "Sunday Pay (£)", "Total Pay (£)",
        "Date Range", "Extracted On", "Source File"
    ]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")

    # Write each row with formulas for pay columns
    for idx, row in df.iterrows():
        excel_row = idx + 2  # Excel is 1-indexed, header is row 1

        name = row["Name"]
        matched = row["Matched As"]
        ratio = row["Ratio"]
        client = row["Client"]
        site_address = row["Site Address"]
        dept = row["Department"]
        wd_hours = row["Weekday Hours"]
        sat_hours = row["Saturday Hours"]
        sun_hours = row["Sunday Hours"]
        rate = row["Rate (£)"]
        date_range = row["Date Range"]
        extracted_on = row["Extracted On"]
        source_file = row["Source File"]

        # Identify column letters
        col_wd = f"G{excel_row}"  # Weekday Hours
        col_sat = f"H{excel_row}"  # Saturday Hours
        col_sun = f"I{excel_row}"  # Sunday Hours
        col_rate = f"J{excel_row}"  # Rate

        # Formulas:
        # Regular Pay = MIN(weekdays, 50) * rate
        reg_formula = f"=MIN({col_wd},50)*{col_rate}"
        # Overtime Pay = MAX(weekdays-50,0)*rate*1.5
        ot_formula = f"=MAX({col_wd}-50,0)*{col_rate}*1.5"
        # Saturday Pay = sat_hours * rate * 1.5
        sat_formula = f"={col_sat}*{col_rate}*1.5"
        # Sunday Pay = sun_hours * rate * 1.75
        sun_formula = f"={col_sun}*{col_rate}*1.75"
        # Total Pay = sum of those four
        tot_formula = f"=K{excel_row}+L{excel_row}+M{excel_row}+N{excel_row}"

        # Append the row
        ws.append([
            name, matched, ratio, client, site_address, dept,
            wd_hours, sat_hours, sun_hours, rate,
            reg_formula, ot_formula, sat_formula, sun_formula, tot_formula,
            date_range, extracted_on, source_file
        ])

    # Adjust column widths (optional)
    for col in ws.columns:
        max_length = 0
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                val_str = str(cell.value)
                if len(val_str) > max_length:
                    max_length = len(val_str)
        adjusted_width = (max_length + 2)
        ws.column_dimensions[col_letter].width = adjusted_width

    wb_out.save(output)
    st.download_button(
        "Download Excel with Formulas",
        data=output.getvalue(),
        file_name="PRL_Timesheet_Report_With_Formulas.xlsx",
        help="This Excel contains formulas for payroll calculations."
    )

else:
    # If no files are uploaded, show a welcome + instructions
    st.markdown("# Welcome to the PRL Timesheet Portal")
    st.markdown(
        """
        **Steps to use:**  
        1. Upload your **Word (.docx)** or **Anel PDF (.pdf)** timesheet files via the sidebar.  
        2. Confirm extracted names vs. pay-detail matches (expand the Debug section).  
        3. If any names are unmatched, use the **Manual Rate Override** in the sidebar.  
        4. Review the final table (read-only).  
        5. Download the final **Excel with Formulas**.  
           - **Regular Pay**: `=MIN(Weekday Hours, 50) * Rate`  
           - **Overtime Pay**: `=MAX(Weekday Hours - 50, 0) * Rate * 1.5`  
           - **Saturday Pay**: `=Saturday Hours * Rate * 1.5`  
           - **Sunday Pay**: `=Sunday Hours * Rate * 1.75`  
           - **Total Pay**: `=[Regular] + [Overtime] + [Saturday] + [Sunday]`  
        6. Send the Excel to the payroll office—all formulas are visible and ready for validation.
        """
    )
    st.info("🚀 Ready when you are! Upload files in the sidebar to get started.")
