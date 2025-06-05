import streamlit as st
import pandas as pd
import docx
import pdfplumber
import re
import unicodedata
from io import BytesIO
from datetime import datetime
from pathlib import Path
from openpyxl import Workbook, load_workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Font, PatternFill
from difflib import SequenceMatcher

# ──────────────────────────────────────────────────────────────────────────────
# PAGE CONFIGURATION
# ──────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="PRL Timesheet Portal",
    page_icon="📑",
    layout="wide"
)

# ──────────────────────────────────────────────────────────────────────────────
# 1) LOAD “pay details.xlsx” FROM ALL SHEETS AND BUILD NAME→RATE MAPPINGS
# ──────────────────────────────────────────────────────────────────────────────

RATE_FILE_PATH = "pay details.xlsx"
DEFAULT_RATE = 15.0
SIMILARITY_THRESHOLD = 0.60

def normalize_name(s: str) -> str:
    s = s.lower().strip()
    s = unicodedata.normalize("NFD", s)
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
    s = re.sub(r"[^a-z0-9\s]", "", s)
    s = re.sub(r"\s+", " ", s)
    return s

@st.cache_data
def load_rate_database(excel_path: str):
    wb = load_workbook(excel_path, read_only=True)
    custom_rates = {}
    normalized_rates = {}
    norm_to_raw = {}

    for sheet in wb.sheetnames:
        df_raw = pd.read_excel(excel_path, sheet_name=sheet, header=None)
        header_row = None
        for idx, val in enumerate(df_raw[0]):
            if isinstance(val, str) and val.strip().lower() == "name":
                second = df_raw.iat[idx, 1]
                if isinstance(second, str) and second.strip().lower() == "pay rate":
                    header_row = idx
                    break
        if header_row is None:
            continue

        df = pd.read_excel(excel_path, sheet_name=sheet, header=header_row)
        if "Name" not in df.columns or "Pay Rate" not in df.columns:
            continue

        df = df[["Name", "Pay Rate"]].copy()
        df["Pay Rate"] = pd.to_numeric(df["Pay Rate"], errors="coerce")
        df = df.dropna(subset=["Name", "Pay Rate"])

        for _, row in df.iterrows():
            raw_name = str(row["Name"]).strip()
            rate = float(row["Pay Rate"])
            custom_rates[raw_name] = rate
            norm = normalize_name(raw_name)
            normalized_rates[norm] = rate
            norm_to_raw[norm] = raw_name

    normalized_keys = list(normalized_rates.keys())
    return custom_rates, normalized_rates, normalized_keys, norm_to_raw

# Sidebar: reload rates button
if st.sidebar.button("🔄 Reload Pay Rates"):
    st.cache_data.clear()
    st.experimental_rerun()

custom_rates, normalized_rates, normalized_keys, norm_to_raw = load_rate_database(RATE_FILE_PATH)

# ──────────────────────────────────────────────────────────────────────────────
# 2) FUZZY LOOKUP AND PAY CALCULATION
# ──────────────────────────────────────────────────────────────────────────────
def lookup_match(name: str):
    norm = normalize_name(name)
    if not norm:
        return None, DEFAULT_RATE, 0.0

    if norm in normalized_rates:
        raw = norm_to_raw[norm]
        return raw, normalized_rates[norm], 1.0

    best_ratio = 0.0
    best_norm = None
    for nk in normalized_keys:
        ratio = SequenceMatcher(None, nk, norm).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_norm = nk

    if best_norm and best_ratio >= SIMILARITY_THRESHOLD:
        raw = norm_to_raw[best_norm]
        return raw, normalized_rates[best_norm], best_ratio

    return None, DEFAULT_RATE, best_ratio

def hhmm_to_float(hhmm: str) -> float:
    try:
        h, m = hhmm.strip().split(":")
        return int(h) + int(m)/60.0
    except:
        return 0.0

def calculate_pay(name: str, daily_data: list[dict]):
    """
    - Mon–Fri total up as weekday_hours.
    - First 50 weekday_hours at rate; any beyond 50 at 1.5×.
    - Saturday at 1.5×, Sunday at 1.75×.
    Returns: (weekday_hours+sat+sun total_raw, rate, total_pay, matched_raw, ratio)
    """
    matched_raw, rate, ratio = lookup_match(name)

    weekday_hours = 0.0
    sat_hours = 0.0
    sun_hours = 0.0

    for entry in daily_data:
        h = entry["hours"]
        wd = entry["weekday"]
        if wd == "Saturday":
            sat_hours += h
        elif wd == "Sunday":
            sun_hours += h
        else:
            weekday_hours += h

    overtime = max(0.0, weekday_hours - 50.0)
    regular_wd = weekday_hours - overtime

    pay_regular = regular_wd * rate
    pay_overtime = overtime * rate * 1.5
    pay_sat = sat_hours * rate * 1.5
    pay_sun = sun_hours * rate * 1.75

    total_pay = pay_regular + pay_overtime + pay_sat + pay_sun
    total_raw = weekday_hours + sat_hours + sun_hours

    return total_raw, rate, total_pay, matched_raw, ratio

# ──────────────────────────────────────────────────────────────────────────────
# 3) DOCX EXTRACTION
# ──────────────────────────────────────────────────────────────────────────────
def extract_timesheet_data(file) -> dict:
    doc = docx.Document(file)
    name = client = site_address = ""
    date_list = []
    daily_data = []

    # 3.1) Find “Client” cell → client & next ALL‐CAPS line as Name
    for table in doc.tables:
        found_c = False
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text or ""
                if "Client" in txt:
                    found_c = True
                    lines = [ln.strip() for ln in txt.splitlines() if ln.strip()]
                    for idx, ln in enumerate(lines):
                        if ln.lower().startswith("client"):
                            parts = re.split(r"Client[:\-\s]+", ln, flags=re.IGNORECASE)
                            if len(parts) > 1:
                                client = parts[1].strip()
                            if idx+1 < len(lines):
                                cand = lines[idx+1]
                                if cand == cand.upper() and len(cand.split()) >= 2:
                                    name = cand.title()
                            break
                    break
            if found_c:
                break
        if found_c:
            break

    # 3.2) Gather daily hours & “Site Address”
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 5:
                hrs_txt = cells[4].text.strip()
                day_txt = cells[1].text.strip()
                date_txt = cells[0].text.strip()
                if hrs_txt and hrs_txt not in ["-", "—", "–"] and day_txt:
                    try:
                        val = float(hrs_txt)
                        daily_data.append({"weekday": day_txt, "hours": val})
                    except:
                        pass
                if re.match(r"\d{2}\.\d{2}\.\d{4}", date_txt):
                    try:
                        d_obj = datetime.strptime(date_txt, "%d.%m.%Y")
                        date_list.append(d_obj)
                    except:
                        pass
            for cell in cells:
                txt = (cell.text or "").strip()
                if "Site Address" in txt and not site_address:
                    m = re.search(r"Site Address[:\-\s]*(.+)", txt)
                    if m:
                        site_address = m.group(1).strip()

    # 3.3) Fallback 1: reversed paragraphs for ALL‐CAPS
    if not name:
        for para in reversed(doc.paragraphs):
            text = (para.text or "").strip()
            if text and text == text.upper() and len(text.split()) >= 2 and "PRL" not in text:
                name = text.title()
                break

    # 3.4) Fallback 2: filename if still no name
    if not name:
        stem = Path(file.name).stem
        name = stem.replace("_", " ").replace("-", " ").title()

    total_raw, rate, total_pay, matched_raw, ratio = calculate_pay(name, daily_data)

    date_range = ""
    if date_list:
        ds = sorted(date_list)
        date_range = f"{ds[0].strftime('%d.%m.%Y')}–{ds[-1].strftime('%d.%m.%Y')}"

    return {
        "Name": name,
        "Matched As": matched_raw or "No match",
        "Ratio": round(ratio, 2),
        "Client": client,
        "Site Address": site_address,
        "Department": "",
        "Total Hours": total_raw,
        "Rate (£)": rate,
        "Calculated Pay (£)": round(total_pay, 2),
        "Date Range": date_range,
        "Extracted On": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "Source File": file.name
    }

# ──────────────────────────────────────────────────────────────────────────────
# 4) PDF EXTRACTION
# ──────────────────────────────────────────────────────────────────────────────
def extract_timesheet_data_pdf(file) -> list[dict]:
    with pdfplumber.open(file) as pdf:
        page1 = pdf.pages[0]
        raw = page1.extract_text()
    lines = raw.split("\n")

    date_range = ""
    for line in lines:
        if line.startswith("Report Range:"):
            m = re.search(
                r"Report Range:\s*(?:Mon|Tue|Wed|Thu|Fri|Sat|Sun)\s+(\d{2}/\d{2}/\d{2})\s+to\s+(?:Mon|Tue|Wed|Thu|Fri|Sat|Sun)\s+(\d{2}/\d{2}/\d{2})",
                line
            )
            if m:
                def fmt(d: str) -> str:
                    dt = datetime.strptime(d, "%d/%m/%y")
                    return dt.strftime("%d.%m.%Y")
                date_range = f"{fmt(m.group(1))}–{fmt(m.group(2))}"
            break

    header_idx = None
    for idx, line in enumerate(lines):
        if "ID Name Paylink" in line and "Tot Sat-Sun" in line:
            header_idx = idx
            break

    if header_idx is None:
        return [{
            "Name": "",
            "Matched As": "No match",
            "Ratio": 0.0,
            "Client": "",
            "Site Address": "",
            "Department": "",
            "Total Hours": 0.0,
            "Rate (£)": 0.0,
            "Calculated Pay (£)": 0.0,
            "Date Range": date_range,
            "Extracted On": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "Source File": file.name
        }]

    results: list[dict] = []
    time_re = re.compile(r"^\d{1,2}:\d{2}$")

    for raw_line in lines[header_idx+1:]:
        if raw_line.strip().startswith("Grand Totals"):
            break
        line = raw_line.strip()
        if not line:
            continue

        tokens = line.split()
        n = len(tokens)
        if n < 10:
            continue

        block = tokens[-9:]
        if not all(time_re.match(t) for t in block):
            continue

        try:
            dash_idx = tokens.index("-")
        except ValueError:
            continue

        name = " ".join(tokens[1:dash_idx]).title()
        daily_data = []
        for wd, idx_tok in [
            ("Monday", 0), ("Tuesday", 1), ("Wednesday", 2),
            ("Thursday", 3), ("Friday", 4), ("Saturday", 6), ("Sunday", 7)
        ]:
            h = hhmm_to_float(block[idx_tok])
            if h > 0:
                daily_data.append({"weekday": wd, "hours": h})

        total_raw, rate, total_pay, matched_raw, ratio = calculate_pay(name, daily_data)

        comp = ""
        if dash_idx + 2 < n:
            comp = " ".join(tokens[dash_idx+1:dash_idx+3])

        site_addr = ""
        if dash_idx + 5 < n:
            site_addr = " ".join(tokens[dash_idx+3:dash_idx+5])

        dept = ""
        dept_start = dash_idx + 5
        dept_end = n - 9
        if dept_start < dept_end:
            dept = " ".join(tokens[dept_start:dept_end]).rstrip("-")

        results.append({
            "Name": name,
            "Matched As": matched_raw or "No match",
            "Ratio": round(ratio, 2),
            "Client": comp,
            "Site Address": site_addr,
            "Department": dept,
            "Total Hours": total_raw,
            "Rate (£)": rate,
            "Calculated Pay (£)": round(total_pay, 2),
            "Date Range": date_range,
            "Extracted On": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "Source File": file.name
        })

    if not results:
        return [{
            "Name": "",
            "Matched As": "No match",
            "Ratio": 0.0,
            "Client": "",
            "Site Address": "",
            "Department": "",
            "Total Hours": 0.0,
            "Rate (£)": 0.0,
            "Calculated Pay (£)": 0.0,
            "Date Range": date_range,
            "Extracted On": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "Source File": file.name
        }]

    return results

# ──────────────────────────────────────────────────────────────────────────────
# 5) LAYOUT: SIDEBAR FOR UPLOADS & OVERRIDES, MAIN AREA FOR RESULTS
# ──────────────────────────────────────────────────────────────────────────────

# Sidebar header
st.sidebar.title("PRL Timesheet Portal")
st.sidebar.markdown("**Upload your DOCX & PDF timesheets** and review matches.")


uploaded_files = st.sidebar.file_uploader(
    "Select Timesheets (.docx, .pdf)",
    type=["docx", "pdf"],
    accept_multiple_files=True
)

# Manual override (only shows once processing is done)
manual_override_section = st.sidebar.expander("✏️ Manual Rate Override", expanded=False)


# ――― PROCESS UPLOADS ―――
if uploaded_files:
    all_rows: list[dict] = []

    for file in uploaded_files:
        lower = file.name.lower()
        if lower.endswith(".docx"):
            data = extract_timesheet_data(file)
            if not data["Name"]:
                stem = Path(file.name).stem
                data["Name"] = stem.replace("_", " ").replace("-", " ").title()
            all_rows.append(data)

        elif lower.endswith(".pdf"):
            pdf_records = extract_timesheet_data_pdf(file)
            for rec in pdf_records:
                if not rec["Name"]:
                    continue
                all_rows.append(rec)

    df = pd.DataFrame(all_rows)

    # ――― DEBUG / MATCH CROSSCHECK ―――
    st.markdown("## 🔎 Extracted vs. Matched Pay-Detail Entries")
    debug_df = (
        df[["Name", "Matched As", "Ratio", "Rate (£)", "Source File"]]
        .drop_duplicates()
        .reset_index(drop=True)
    )
    # Show debug in a collapsible container
    with st.expander("Show Name-Match Debug Table"):
        st.table(debug_df)

    # List any “No match”
    no_match = debug_df.loc[debug_df["Matched As"] == "No match", "Name"].tolist()
    if no_match:
        st.error(
            "⚠️ The following names did not match any pay-detail entry (ratio < {:.2f}):\n\n{}".format(
                SIMILARITY_THRESHOLD, "\n".join(f"- {nm}" for nm in no_match)
            )
        )

    # ――― MANUAL OVERRIDE ―――
    unmatched = debug_df.loc[debug_df["Matched As"] == "No match", "Name"].tolist()
    if unmatched:
        with manual_override_section:
            st.markdown("**Select an extracted name to map to a pay-detail entry:**")
            override_name = st.selectbox("Unmatched Extracted Name", options=[""] + unmatched)
            if override_name:
                chosen = st.selectbox(
                    "Map to Pay-Detail Contractor",
                    options=[""] + list(custom_rates.keys())
                )
                if chosen and st.button("Apply Override"):
                    correct_rate = custom_rates.get(chosen)
                    df.loc[df["Name"] == override_name, "Matched As"] = chosen
                    df.loc[df["Name"] == override_name, "Rate (£)"] = correct_rate
                    # Recalculate pay for those rows
                    df.loc[df["Name"] == override_name, ["Calculated Pay (£)", "Total Hours"]] = df.loc[df["Name"] == override_name].apply(
                        lambda row: pd.Series(
                            calculate_pay(row["Name"], [{"weekday": wd, "hours": hhmm_to_float(hh)} 
                                for wd, hh in zip(
                                    ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"],
                                    row.get("Raw Hours List", [0,0,0,0,0,0,0])
                                )]
                            )[[2,0]]  # pay, total_hours
                        ),
                        axis=1
                    )
                    st.success(f"✅ Overrode rate for '{override_name}' → £{correct_rate:.2f}")

    # ――― MAIN RESULTS LAYOUT ―――
    st.markdown("---")
    st.markdown("## 📋 Full Processed Timesheets")
    st.dataframe(df)

    # Weekly summary (group by matched name)
    summary_df = (
        df.groupby("Matched As")[["Calculated Pay (£)", "Total Hours"]]
        .sum()
        .reset_index()
        .rename(columns={"Matched As": "Name"})
    )

    st.markdown("## 💰 Weekly Summary")
    col1, col2 = st.columns([2, 1])
    with col1:
        st.dataframe(summary_df)
    with col2:
        total_pay_all = summary_df["Calculated Pay (£)"].sum()
        total_hours_all = summary_df["Total Hours"].sum()
        st.metric(label="Total Hours This Period", value=f"{total_hours_all:.2f}")
        st.metric(label="Total Pay This Period", value=f"£{total_pay_all:.2f}")

    # ――― EXPORT BUTTON ―――
    st.markdown("---")
    st.markdown("### 📥 Download Excel Report")
    output = BytesIO()
    wb_out = Workbook()

    ws1 = wb_out.active
    ws1.title = "Timesheets"
    for r in dataframe_to_rows(df, index=False, header=True):
        ws1.append(r)
    for cell in ws1[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")

    ws2 = wb_out.create_sheet("Weekly Summary")
    for r in dataframe_to_rows(summary_df, index=False, header=True):
        ws2.append(r)
    for cell in ws2[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CFCFCF", end_color="CFCFCF", fill_type="solid")

    wb_out.save(output)
    st.download_button(
        "Download PRL_Timesheet_Report.xlsx",
        data=output.getvalue(),
        file_name="PRL_Timesheet_Report.xlsx",
        help="Click to download the consolidated report."
    )

else:
    # If no files uploaded yet, show a friendly welcome message and instructions
    st.markdown("# Welcome to the PRL Timesheet Portal")
    st.markdown(
        """
        1. Use the sidebar to upload your **Word (.docx) or Anel PDF (.pdf)** timesheet files.
        2. The portal will automatically extract hours, match each contractor to their day rate,
           apply overtime/weekend multipliers, and display the results below.
        3. You can review the “Debug” table (collapsed by default) to confirm how each name was matched.
        4. If any names aren’t found in **pay details.xlsx**, use the **Manual Rate Override** section in the sidebar.
        5. Once everything looks correct, download your final Excel report (Timesheets + Weekly Summary).
        """
    )
    st.info("🚀 Ready when you are! Upload files via the sidebar to begin.")

