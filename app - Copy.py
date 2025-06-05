import streamlit as st
import pandas as pd
import docx
import pdfplumber
import re
import unicodedata
from io import BytesIO
from datetime import datetime
from pathlib import Path
from openpyxl import Workbook, load_workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Font, PatternFill
from difflib import SequenceMatcher

# ──────────────────────────────────────────────────────────────────────────────
# 1) Load “pay details.xlsx” from ALL SHEETS, build normalized {name → rate}
# ──────────────────────────────────────────────────────────────────────────────

RATE_FILE_PATH = "pay details.xlsx"
DEFAULT_RATE = 15.0
SIMILARITY_THRESHOLD = 0.60  # Only accept fuzzy matches if ratio ≥ this

def normalize_name(s: str) -> str:
    """
    Lowercase, strip accents, remove punctuation, collapse spaces.
    """
    s = s.lower().strip()
    s = unicodedata.normalize("NFD", s)
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
    s = re.sub(r"[^a-z0-9\s]", "", s)
    s = re.sub(r"\s+", " ", s)
    return s

@st.cache_data
def load_rate_database(excel_path: str):
    """
    Iterates all sheets in pay details.xlsx. In each sheet:
      1) Find the header row where A="Name" and B="Pay Rate".
      2) Read that sheet with header=that_row.
      3) Keep only rows where "Name" and numeric "Pay Rate" are present.
      4) Build:
         • custom_rates: {raw_name → day_rate}
         • normalized_rates: {normalize_name(raw_name) → day_rate}
         • norm_to_raw:      {normalize_name(raw_name) → raw_name}
      • Collect all normalized_keys for fuzzy matching.
    """
    wb = load_workbook(excel_path, read_only=True)
    custom_rates = {}
    normalized_rates = {}
    norm_to_raw = {}

    for sheet in wb.sheetnames:
        df_raw = pd.read_excel(excel_path, sheet_name=sheet, header=None)
        header_row = None
        for idx, val in enumerate(df_raw[0]):
            if isinstance(val, str) and val.strip().lower() == "name":
                second = df_raw.iat[idx, 1]
                if isinstance(second, str) and second.strip().lower() == "pay rate":
                    header_row = idx
                    break
        if header_row is None:
            continue

        df = pd.read_excel(excel_path, sheet_name=sheet, header=header_row)
        if "Name" not in df.columns or "Pay Rate" not in df.columns:
            continue

        df = df[["Name", "Pay Rate"]].copy()
        df["Pay Rate"] = pd.to_numeric(df["Pay Rate"], errors="coerce")
        df = df.dropna(subset=["Name", "Pay Rate"])

        for _, row in df.iterrows():
            raw_name = str(row["Name"]).strip()
            rate = float(row["Pay Rate"])
            custom_rates[raw_name] = rate
            norm = normalize_name(raw_name)
            normalized_rates[norm] = rate
            norm_to_raw[norm] = raw_name

    normalized_keys = list(normalized_rates.keys())
    return custom_rates, normalized_rates, normalized_keys, norm_to_raw

# Allow reloading the Excel at any time
if st.button("🔄 Reload rates"):
    st.cache_data.clear()
    st.experimental_rerun()

custom_rates, normalized_rates, normalized_keys, norm_to_raw = load_rate_database(RATE_FILE_PATH)

# ──────────────────────────────────────────────────────────────────────────────
# 2) SequenceMatcher‐based fuzzy lookup
# ──────────────────────────────────────────────────────────────────────────────
def lookup_match(name: str):
    """
    Given an extracted 'name':
      1) Normalize it.
      2) If normalize(name) in normalized_rates → exact match; return (raw_name, rate, 1.0).
      3) Otherwise, iterate all normalized_keys, compute SequenceMatcher ratio:
         best_ratio = max( ratio(norm_key, norm_name) for norm_key in normalized_keys ).
         If best_ratio ≥ SIMILARITY_THRESHOLD, return (raw_name_for_best_key, rate_for_best_key, best_ratio).
      4) Else return (None, DEFAULT_RATE, best_ratio).
    """
    norm = normalize_name(name)
    if not norm:
        return None, DEFAULT_RATE, 0.0

    # Exact match
    if norm in normalized_rates:
        raw = norm_to_raw[norm]
        return raw, normalized_rates[norm], 1.0

    # Fuzzy via SequenceMatcher
    best_ratio = 0.0
    best_norm = None
    for nk in normalized_keys:
        ratio = SequenceMatcher(None, nk, norm).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_norm = nk

    if best_norm and best_ratio >= SIMILARITY_THRESHOLD:
        raw = norm_to_raw[best_norm]
        return raw, normalized_rates[best_norm], best_ratio

    return None, DEFAULT_RATE, best_ratio

# ──────────────────────────────────────────────────────────────────────────────
# 3) Convert HH:MM to float
# ──────────────────────────────────────────────────────────────────────────────
def hhmm_to_float(hhmm: str) -> float:
    try:
        h, m = hhmm.strip().split(":")
        return int(h) + int(m) / 60.0
    except:
        return 0.0

# ──────────────────────────────────────────────────────────────────────────────
# 4) Calculate pay with overtime & weekend rules
# ──────────────────────────────────────────────────────────────────────────────
def calculate_pay(name: str, daily_data: list[dict]):
    """
    - Sum raw hours for Mon-Fri, Sat, Sun separately.
    - Weekday overtime: first 50 Mon-Fri hours at rate; any beyond 50 at 1.5×.
    - Saturday hours at 1.5×, Sunday hours at 1.75× regardless of overtime.
    Returns: (total_raw_hours, rate, pay_amount, matched_raw_name, ratio)
    """
    # Lookup matched name and rate
    matched_raw, rate, ratio = lookup_match(name)

    # Tally hours by day-type
    weekday_hours = 0.0
    saturday_hours = 0.0
    sunday_hours = 0.0

    for entry in daily_data:
        h = entry["hours"]
        wd = entry["weekday"]
        if wd == "Saturday":
            saturday_hours += h
        elif wd == "Sunday":
            sunday_hours += h
        else:
            # Monday–Friday (treated as weekday for overtime)
            weekday_hours += h

    # Compute overtime on weekday hours only
    overtime_hours = max(0.0, weekday_hours - 50.0)
    regular_weekday_hours = weekday_hours - overtime_hours

    # Compute pay segments
    pay_regular = regular_weekday_hours * rate
    pay_overtime = overtime_hours * rate * 1.5
    pay_sat = saturday_hours * rate * 1.5
    pay_sun = sunday_hours * rate * 1.75

    total_pay = pay_regular + pay_overtime + pay_sat + pay_sun
    total_raw = weekday_hours + saturday_hours + sunday_hours

    return total_raw, rate, total_pay, matched_raw, ratio

# ──────────────────────────────────────────────────────────────────────────────
# 5) DOCX extractor
# ──────────────────────────────────────────────────────────────────────────────
def extract_timesheet_data(file) -> dict:
    """
    Parse a .docx:
      - Scan table cells for “Client” and extract the next ALL‐CAPS line as Name.
      - Collect daily hours from column index 4 (if numeric) keyed by weekday in index 1.
      - Capture “Site Address: …” if present.
      - Fallback 1: reversed paragraphs → first ALL‐CAPS line as Name.
      - Fallback 2: filename → Name.
      - Call calculate_pay(...) to get raw hours, rate, pay, matched, ratio.
      - Build a dict with all fields.
    """
    doc = docx.Document(file)
    name = client = site_address = ""
    date_list = []
    daily_data = []

    # 5.1) Find “Client” cell → get client + name beneath
    for table in doc.tables:
        client_found = False
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text or ""
                if "Client" in txt:
                    client_found = True
                    lines = [ln.strip() for ln in txt.splitlines() if ln.strip()]
                    for idx, ln in enumerate(lines):
                        if ln.lower().startswith("client"):
                            parts = re.split(r"Client[:\-\s]+", ln, flags=re.IGNORECASE)
                            if len(parts) > 1:
                                client = parts[1].strip()
                            # Next line ALL‐CAPS → Name
                            if idx + 1 < len(lines):
                                cand = lines[idx + 1]
                                if cand == cand.upper() and len(cand.split()) >= 2:
                                    name = cand.title()
                            break
                    break
            if client_found:
                break
        if client_found:
            break

    # 5.2) Collect daily hours & site address
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 5:
                hrs_txt = cells[4].text.strip()
                day_txt = cells[1].text.strip()
                date_txt = cells[0].text.strip()
                if hrs_txt and hrs_txt not in ["-", "–", "—"] and day_txt:
                    try:
                        val = float(hrs_txt)
                        daily_data.append({"weekday": day_txt, "hours": val})
                    except:
                        pass
                if re.match(r"\d{2}\.\d{2}\.\d{4}", date_txt):
                    try:
                        d_obj = datetime.strptime(date_txt, "%d.%m.%Y")
                        date_list.append(d_obj)
                    except:
                        pass
            for cell in cells:
                txt = (cell.text or "").strip()
                if "Site Address" in txt and not site_address:
                    m = re.search(r"Site Address[:\-\s]*(.+)", txt)
                    if m:
                        site_address = m.group(1).strip()

    # 5.3) Fallback 1: reversed paragraphs for ALL‐CAPS name
    if not name:
        for para in reversed(doc.paragraphs):
            text = (para.text or "").strip()
            if text and text == text.upper() and len(text.split()) >= 2 and "PRL" not in text:
                name = text.title()
                break

    # 5.4) Fallback 2: filename if still no name
    if not name:
        stem = Path(file.name).stem
        name = stem.replace("_", " ").replace("-", " ").title()

    total_raw, rate, total_pay, matched_raw, ratio = calculate_pay(name, daily_data)

    date_range = ""
    if date_list:
        ds = sorted(date_list)
        date_range = f"{ds[0].strftime('%d.%m.%Y')}–{ds[-1].strftime('%d.%m.%Y')}"

    return {
        "Name": name,
        "Matched As": matched_raw or "No match",
        "Ratio": round(ratio, 2),
        "Client": client,
        "Site Address": site_address,
        "Department": "",
        "Total Hours": total_raw,
        "Rate (£)": rate,
        "Calculated Pay (£)": round(total_pay, 2),
        "Date Range": date_range,
        "Extracted On": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "Source File": file.name
    }

# ──────────────────────────────────────────────────────────────────────────────
# 6) PDF extractor (include “Matched As” & “Ratio” as well)
# ──────────────────────────────────────────────────────────────────────────────
def extract_timesheet_data_pdf(file) -> list[dict]:
    """
    Parse an Anel PDF:
      - Extract “Report Range: … to …” for date_range.
      - Find header “ID Name Paylink” + “Tot Sat-Sun”.
      - For each line until “Grand Totals”:
        • tokens = line.split()
        • last 9 tokens are H:MM block → daily_data
        • dash_idx = tokens.index("-")
        • Name = tokens[1:dash_idx]
        • Company, Site, Dept parsed
        • Call calculate_pay(name, daily_data)
        • Append dict with fields including “Matched As” & “Ratio”.
    """
    with pdfplumber.open(file) as pdf:
        page1 = pdf.pages[0]
        raw = page1.extract_text()
    lines = raw.split("\n")

    date_range = ""
    for line in lines:
        if line.startswith("Report Range:"):
            m = re.search(
                r"Report Range:\s*(?:Mon|Tue|Wed|Thu|Fri|Sat|Sun)\s+(\d{2}/\d{2}/\d{2})\s+to\s+(?:Mon|Tue|Wed|Thu|Fri|Sat|Sun)\s+(\d{2}/\d{2}/\d{2})",
                line,
            )
            if m:
                def fmt(d: str) -> str:
                    dt = datetime.strptime(d, "%d/%m/%y")
                    return dt.strftime("%d.%m.%Y")
                date_range = f"{fmt(m.group(1))}–{fmt(m.group(2))}"
            break

    header_idx = None
    for idx, line in enumerate(lines):
        if "ID Name Paylink" in line and "Tot Sat-Sun" in line:
            header_idx = idx
            break

    if header_idx is None:
        return [{
            "Name": "",
            "Matched As": "No match",
            "Ratio": 0.0,
            "Client": "",
            "Site Address": "",
            "Department": "",
            "Total Hours": 0.0,
            "Rate (£)": 0.0,
            "Calculated Pay (£)": 0.0,
            "Date Range": date_range,
            "Extracted On": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "Source File": file.name
        }]

    results: list[dict] = []
    time_re = re.compile(r"^\d{1,2}:\d{2}$")

    for raw_line in lines[header_idx + 1 :]:
        if raw_line.strip().startswith("Grand Totals"):
            break
        line = raw_line.strip()
        if not line:
            continue

        tokens = line.split()
        n = len(tokens)
        if n < 10:
            continue

        block = tokens[-9:]
        if not all(time_re.match(t) for t in block):
            continue

        try:
            dash_idx = tokens.index("-")
        except ValueError:
            continue

        name = " ".join(tokens[1:dash_idx]).title()
        # Build daily_data
        daily_data = []
        for wd, idx_tok in [
            ("Monday", 0),
            ("Tuesday", 1),
            ("Wednesday", 2),
            ("Thursday", 3),
            ("Friday", 4),
            ("Saturday", 6),
            ("Sunday", 7),
        ]:
            h = hhmm_to_float(block[idx_tok])
            if h > 0:
                daily_data.append({"weekday": wd, "hours": h})

        total_raw, rate, total_pay, matched_raw, ratio = calculate_pay(name, daily_data)

        comp = ""
        if dash_idx + 2 < n:
            comp = " ".join(tokens[dash_idx + 1 : dash_idx + 3])

        site_addr = ""
        if dash_idx + 5 < n:
            site_addr = " ".join(tokens[dash_idx + 3 : dash_idx + 5])

        dept = ""
        dept_start = dash_idx + 5
        dept_end = n - 9
        if dept_start < dept_end:
            dept = " ".join(tokens[dept_start : dept_end]).rstrip("-")

        results.append({
            "Name": name,
            "Matched As": matched_raw or "No match",
            "Ratio": round(ratio, 2),
            "Client": comp,
            "Site Address": site_addr,
            "Department": dept,
            "Total Hours": total_raw,
            "Rate (£)": rate,
            "Calculated Pay (£)": round(total_pay, 2),
            "Date Range": date_range,
            "Extracted On": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "Source File": file.name
        })

    if not results:
        return [{
            "Name": "",
            "Matched As": "No match",
            "Ratio": 0.0,
            "Client": "",
            "Site Address": "",
            "Department": "",
            "Total Hours": 0.0,
            "Rate (£)": 0.0,
            "Calculated Pay (£)": 0.0,
            "Date Range": date_range,
            "Extracted On": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "Source File": file.name
        }]

    return results

# ──────────────────────────────────────────────────────────────────────────────
# 7) Streamlit UI: show debug table + main tables + export
# ──────────────────────────────────────────────────────────────────────────────

st.title("📑 PRL Timesheet Aggregator (Overtime & Weekend Handling)")

uploaded_files = st.file_uploader(
    "Upload Word (.docx) or Anel PDF (.pdf) Timesheets",
    type=["docx", "pdf"],
    accept_multiple_files=True,
)

if uploaded_files:
    rows: list[dict] = []

    # 7.1) Extract each file (DOCX or PDF)
    for file in uploaded_files:
        lower = file.name.lower()
        if lower.endswith(".docx"):
            data = extract_timesheet_data(file)
            if not data["Name"]:
                # Fallback: filename → Name
                stem = Path(file.name).stem
                data["Name"] = stem.replace("_", " ").replace("-", " ").title()
            rows.append(data)

        elif lower.endswith(".pdf"):
            pdf_records = extract_timesheet_data_pdf(file)
            for entry in pdf_records:
                if not entry["Name"]:
                    continue
                rows.append(entry)

    df = pd.DataFrame(rows)

    # 7.2) Show a debug snapshot: extracted name → matched pay‐detail name → ratio → rate
    st.markdown("### 🔎 Debug: Extracted vs. Matched Pay‐Detail Entries")
    debug_df = (
        df[["Name", "Matched As", "Ratio", "Rate (£)", "Source File"]]
        .drop_duplicates()
        .reset_index(drop=True)
    )
    st.dataframe(debug_df)

    # 7.3) List all “No match” cases
    no_match = debug_df.loc[debug_df["Matched As"] == "No match", "Name"].tolist()
    if no_match:
        st.warning(
            "These extracted names did not match any pay‐detail entry (ratio < "
            f"{SIMILARITY_THRESHOLD:.2f}):\n\n"
            + "\n".join(f"- {nm}" for nm in no_match)
        )

    # 7.4) Show the full processed timesheet table
    st.markdown("### 📋 Full Processed Timesheets")
    st.dataframe(df)

    # 7.5) Show Weekly Summary (group by “Matched As”)
    summary_df = (
        df.groupby("Matched As")[["Calculated Pay (£)", "Total Hours"]]
        .sum()
        .reset_index()
        .rename(columns={"Matched As": "Name"})
    )
    st.markdown("### 💰 Weekly Summary")
    st.dataframe(summary_df)

    # 7.6) Export to Excel
    output = BytesIO()
    wb_out = Workbook()

    ws1 = wb_out.active
    ws1.title = "Timesheets"
    for r in dataframe_to_rows(df, index=False, header=True):
        ws1.append(r)
    for cell in ws1[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(
            start_color="D9D9D9", end_color="D9D9D9", fill_type="solid"
        )

    ws2 = wb_out.create_sheet("Weekly Summary")
    for r in dataframe_to_rows(summary_df, index=False, header=True):
        ws2.append(r)
    for cell in ws2[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(
            start_color="CFCFCF", end_color="CFCFCF", fill_type="solid"
        )

    wb_out.save(output)
    st.download_button(
        "📥 Download Excel Report",
        data=output.getvalue(),
        file_name="PRL_Timesheet_Report.xlsx",
    )
