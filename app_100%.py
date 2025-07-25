import os
import re
import unicodedata
import zipfile
from pathlib import Path
from io import BytesIO
from datetime import datetime, date, timedelta

import streamlit as st
import pandas as pd
import docx
import pdfplumber
from openpyxl import load_workbook

# ==== DB Connection (Postgres vs SQLite) ====
IS_PG = "DATABASE_URL" in os.environ
if IS_PG:
    import psycopg2
    from urllib.parse import urlparse
    url = urlparse(os.environ["DATABASE_URL"])
    conn = psycopg2.connect(
        dbname=url.path[1:], user=url.username,
        password=url.password, host=url.hostname, port=url.port
    )
else:
    import sqlite3
    conn = sqlite3.connect("timesheets.db", check_same_thread=False)

c = conn.cursor()
# Add is_paid flag
c.execute("""
CREATE TABLE IF NOT EXISTS timesheet_entries (
    id SERIAL PRIMARY KEY,
    name TEXT,
    matched_as TEXT,
    ratio REAL,
    client TEXT,
    site_address TEXT,
    department TEXT,
    weekday_hours REAL,
    saturday_hours REAL,
    sunday_hours REAL,
    rate REAL,
    date_range TEXT,
    extracted_on TEXT,
    source_file TEXT,
    upload_timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    is_paid BOOLEAN DEFAULT FALSE
)
""")
conn.commit()

# ==== Helpers ====
def normalize_name(name: str) -> str:
    nfkd = unicodedata.normalize("NFKD", name)
    only_ascii = nfkd.encode("ASCII", "ignore").decode("utf-8")
    return re.sub(r"[^a-zA-Z]", "", only_ascii).lower()

def load_rate_database(source):
    custom, normed, to_raw = {}, {}, {}
    wb = load_workbook(source, data_only=True)
    for sheet in wb.sheetnames:
        df0 = pd.read_excel(source, sheet_name=sheet, header=None)
        header_row = None
        for i,v in enumerate(df0.iloc[:,0]):
            if (
                isinstance(v,str) and v.strip().lower()=="name"
                and isinstance(df0.iat[i,1],str)
                and df0.iat[i,1].strip().lower()=="pay rate"
            ):
                header_row=i; break
        if header_row is None: continue
        df = pd.read_excel(source, sheet_name=sheet, header=header_row)
        if "Name" not in df or "Pay Rate" not in df: continue
        df = df[["Name","Pay Rate"]].dropna()
        df["Pay Rate"] = pd.to_numeric(df["Pay Rate"], errors="coerce")
        for _,r in df.iterrows():
            raw=str(r["Name"]).strip()
            rate=float(r["Pay Rate"])
            custom[raw]=rate
            n=normalize_name(raw)
            normed[n]=rate
            to_raw[n]=raw
    return custom,normed,to_raw

def lookup_match(name: str):
    n=normalize_name(name)
    if n in normalized_rates:
        return norm_to_raw[n], normalized_rates[n], 1.0
    return name, 15.0, 0.0

# ==== Rate‑Loader (multi‑sheet) ====
RATE_FILE="pay_rates.xlsx"
uploads=st.sidebar.file_uploader(
    "➕ Upload pay‑rate XLSX(s)",
    type=["xlsx"], accept_multiple_files=True
)
custom_rates,normalized_rates,norm_to_raw={},{} ,{}
def _merge(src):
    cr,nr,nt=load_rate_database(src)
    custom_rates.update(cr)
    normalized_rates.update(nr)
    norm_to_raw.update(nt)

if uploads:
    for f in uploads: _merge(f)
    st.sidebar.success(f"Merged {len(uploads)} rate sheet(s).")
elif Path(RATE_FILE).exists():
    try:
        _merge(RATE_FILE)
        st.sidebar.info(f"Loaded local `{RATE_FILE}`.")
    except Exception as e:
        st.sidebar.error(f"Error loading `{RATE_FILE}`: {e}")
if not normalized_rates:
    st.sidebar.warning("No pay‑rates found; defaulting to £15/hr.")

# ==== Extraction Logic (unchanged) ====
def extract_from_docx(file) -> list[dict]:
    doc=docx.Document(file)
    if not doc.tables: return []
    tbl=doc.tables[0]
    hdr=tbl.rows[0].cells[0].text.split("\n")
    hdr=[h.strip() for h in hdr if h.strip()]
    client=name=site=None
    for i,line in enumerate(hdr):
        low=line.lower()
        if low.startswith("client"):
            client=line.split(None,1)[1].strip()
            if i+1<len(hdr): name=hdr[i+1].strip()
        if low.startswith("site address"):
            parts=line.split("\t",1)
            site=parts[1].strip() if len(parts)>1 else None
    # find "Date" row
    header_row=next((i for i,row in enumerate(tbl.rows)
                     if row.cells[0].text.strip().lower()=="date"), None)
    if header_row is None: return []
    date_re=re.compile(r"\d{2}\.\d{2}\.\d{4}")
    wd=sa=su=0.0; dates=[]
    for row in tbl.rows[header_row+1:]:
        dt=row.cells[0].text.strip()
        if not date_re.match(dt): continue
        dates.append(dt)
        day=row.cells[1].text.strip().lower()
        try: hrs=float(row.cells[4].text.strip())
        except: hrs=0.0
        if day=="saturday": sa+=hrs
        elif day=="sunday": su+=hrs
        else: wd+=hrs
    if not dates: return []
    dr=f"{min(dates)}–{max(dates)}"
    matched,rate,ratio=lookup_match(name or "")
    return [{
        "id": None,                # placeholder
        "name":name or "",
        "matched_as":matched,
        "ratio":ratio,
        "client":client or "",
        "site_address":site or "",
        "department":"",
        "weekday_hours":wd,
        "saturday_hours":sa,
        "sunday_hours":su,
        "rate":rate,
        "date_range":dr,
        "extracted_on":datetime.now().isoformat(),
        "source_file":None,        # set below
        "is_paid":False
    }]

def extract_from_pdf(file)->list[dict]:
    return []

# ==== Sidebar Uploader ====
st.sidebar.header("Upload Timesheets")
st.sidebar.markdown("""
1. Upload **.docx**, **.pdf** or **.zip**  
2. Check duplicates & mark paid  
3. Download Excel  
""")
uploaded=st.sidebar.file_uploader(
    "Choose timesheet file(s)", accept_multiple_files=True
)

# ==== Main Tabs ====
tabs=st.tabs([
    "Upload & Review",
    "History",
    "Matches",
    "Dashboard",
    "Settings"
])

# ---- 1) Upload & Review ----
with tabs[0]:
    st.header("📤 Upload & Review Timesheets")
    if not uploaded:
        st.info("Upload .docx/.pdf/.zip to begin.")
    else:
        total=len(uploaded)
        progress=st.progress(0)
        summaries=[]

        with st.spinner(f"Processing {total} file(s)…"):
            for i,uf in enumerate(uploaded):
                st.write(f"➡️ {uf.name} ({i+1}/{total})")
                lower=uf.name.lower()
                def handle_recs(recs):
                    for r in recs:
                        r["source_file"]=uf.name
                        summaries.append(r)

                if lower.endswith(".zip"):
                    try:
                        z=zipfile.ZipFile(uf)
                        members=[m for m in z.namelist()
                                 if m.lower().endswith((".docx",".pdf"))]
                        if not members:
                            st.warning(f"{uf.name} empty ZIP.")
                        for m in members:
                            data=z.read(m)
                            buf=BytesIO(data); buf.name=m
                            recs=(extract_from_docx(buf)
                                  if m.lower().endswith(".docx")
                                  else extract_from_pdf(buf))
                            st.write(f" • {m}: {len(recs)} rec(s)")
                            handle_recs(recs)
                    except zipfile.BadZipFile:
                        st.error(f"{uf.name} invalid ZIP.")
                elif lower.endswith(".docx"):
                    st.write(" • DOCX…")
                    recs=extract_from_docx(uf)
                    st.write(f"   → {len(recs)} rec(s)")
                    handle_recs(recs)
                elif lower.endswith(".pdf"):
                    st.write(" • PDF…")
                    recs=extract_from_pdf(uf)
                    st.write(f"   → {len(recs)} rec(s)")
                    handle_recs(recs)
                else:
                    st.warning(f"Unsupported: {uf.name}")

                progress.progress((i+1)/total)

        if not summaries:
            st.error("No records extracted.")
        else:
            df=pd.DataFrame(summaries)
            with st.expander("🔍 Raw summaries"):
                st.dataframe(df.drop(columns=["id"]), use_container_width=True)

            # Detect duplicates
            existing,new=[],[]
            for r in summaries:
                if IS_PG:
                    c.execute(
                        "SELECT COUNT(*) FROM timesheet_entries WHERE name=%s AND date_range=%s",
                        (r["name"],r["date_range"])
                    )
                else:
                    c.execute(
                        "SELECT COUNT(*) FROM timesheet_entries WHERE name=? AND date_range=?",
                        (r["name"],r["date_range"])
                    )
                (cnt,)=c.fetchone()
                (existing if cnt>0 else new).append(r)

            if existing:
                st.warning("⚠️ Duplicates (skipped):")
                st.dataframe(
                    pd.DataFrame(existing)[["name","date_range","source_file"]],
                    use_container_width=True
                )

            # Excel export for all
            towrite=BytesIO()
            with pd.ExcelWriter(towrite, engine="openpyxl") as writer:
                df.drop(columns=["id"]).to_excel(writer, index=False)
            towrite.seek(0)
            st.download_button(
                "📥 Download All Summaries",
                data=towrite,
                file_name=f"summaries_{date.today()}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

            # Persist new
            if new:
                ph=",".join("%s" if IS_PG else "?" for _ in range(13))
                sql=f"""
                INSERT INTO timesheet_entries
                  (name,matched_as,ratio,client,site_address,department,
                   weekday_hours,saturday_hours,sunday_hours,rate,
                   date_range,extracted_on,source_file)
                VALUES({ph})
                """
                for r in new:
                    params=[
                        r["name"],r["matched_as"],r["ratio"],
                        r["client"],r["site_address"],r["department"],
                        r["weekday_hours"],r["saturday_hours"],
                        r["sunday_hours"],r["rate"],
                        r["date_range"],r["extracted_on"],r["source_file"]
                    ]
                    c.execute(sql, params)
                conn.commit()
                st.success(f"Inserted {len(new)} new rec(s).")
            else:
                st.info("No new records to insert.")

# ---- 2) History ----
with tabs[1]:
    st.header("🗃️ Upload History")
    st.markdown("Quick presets or custom range:")
    # presets
    presets=["Last 30 days","This month","Last month","Year to date","Custom"]
    choice=st.selectbox("Date range",presets)
    today=date.today()
    if choice!="Custom":
        if choice=="Last 30 days":
            start,end=today-timedelta(30),today
        elif choice=="This month":
            start,end=today.replace(day=1),today
        elif choice=="Last month":
            first=today.replace(day=1)
            last=first-timedelta(1)
            start,end=last.replace(day=1),last
        elif choice=="Year to date":
            start,end=date(today.year,1,1),today
        st.write(f"Showing **{choice}**: {start} → {end}")
    else:
        start,end=st.date_input(
            "Custom range",
            value=(today-timedelta(30),today),
            min_value=date(2020,1,1),max_value=today
        )
    # fetch
    c.execute("""
        SELECT id,name,matched_as,ratio,client,site_address,department,
               weekday_hours,saturday_hours,sunday_hours,rate AS rate,
               date_range,extracted_on,source_file,upload_timestamp,is_paid
        FROM timesheet_entries
        ORDER BY upload_timestamp DESC
    """)
    hist=pd.DataFrame(c.fetchall(),columns=[
        "id","Name","Matched As","Ratio","Client","Site Address",
        "Department","Weekday Hours","Saturday Hours","Sunday Hours",
        "Rate (£)","Date Range","Extracted On","Source File",
        "Upload Timestamp","Paid?"
    ])
    hist["Upload Timestamp"]=pd.to_datetime(hist["Upload Timestamp"]).dt.date
    mask=(hist["Upload Timestamp"]>=start)&(hist["Upload Timestamp"]<=end)
    view=hist.loc[mask].reset_index(drop=True)
    if view.empty:
        st.info("No entries in this range.")
    else:
        st.markdown("### 📅 Weekly Summary")
        summary=(view.groupby("Date Range").agg(
            Entries=("Name","count"),
            Weekday_Hours=("Weekday Hours","sum"),
            Sat_Hours=("Saturday Hours","sum"),
            Sun_Hours=("Sunday Hours","sum"),
            Total_Pay=(lambda df:(df["Weekday Hours"]+df["Saturday Hours"]+df["Sunday Hours"])*df["Rate (£)"])
        ).reset_index().sort_values("Date Range",ascending=False))
        st.dataframe(summary,use_container_width=True)

        # bulk actions via data_editor
        st.markdown("### ✅ Bulk Actions")
        selected=st.data_editor(
            view,
            column_config={"id":{"hidden":True},"Paid?":{"type":"boolean"}},
            row_selectable="multi",
            use_container_width=True
        )
        sel_ids=selected["id"].tolist()
        if sel_ids:
            col1,col2,col3=st.columns(3)
            if col1.button("Delete selected"):
                ph=",".join("%s" if IS_PG else "?" for _ in sel_ids)
                c.execute(f"DELETE FROM timesheet_entries WHERE id IN ({ph})", sel_ids)
                conn.commit()
                st.success(f"Deleted {len(sel_ids)} record(s).")
            if col2.button("Export selected"):
                df_sel=view[view["id"].isin(sel_ids)]
                buf=BytesIO()
                with pd.ExcelWriter(buf,engine="openpyxl") as w: df_sel.to_excel(w,index=False)
                buf.seek(0)
                st.download_button( "Download Export", data=buf,
                    file_name=f"export_{date.today()}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            if col3.button("Mark paid"):
                ph=",".join("%s" if IS_PG else "?" for _ in sel_ids)
                c.execute(f"UPDATE timesheet_entries SET is_paid=TRUE WHERE id IN ({ph})", sel_ids)
                conn.commit()
                st.success(f"Marked {len(sel_ids)} paid.")

# ---- 3) Matches ----
with tabs[2]:
    st.header("🔗 Name → Rate Matches")
    st.markdown("Inline‑edit matched name or confidence:")
    c.execute("SELECT DISTINCT name,matched_as,ratio FROM timesheet_entries ORDER BY name")
    dfm=pd.DataFrame(c.fetchall(),columns=["Timesheet Name","Matched Rate Name","Confidence"])
    if dfm.empty:
        st.info("No matches yet.")
    else:
        edited=st.data_editor(
            dfm,
            column_config={
                "Timesheet Name":{"disabled":True},
                "Matched Rate Name":{},
                "Confidence":{"min_value":0.0,"max_value":1.0}
            },
            row_selectable=False,
            use_container_width=True
        )
        if st.button("Save match edits"):
            diffs=edited.merge(dfm,indicator=True,how="outer").loc[lambda d: d["_merge"]!="both"]
            for _,r in diffs.iterrows():
                name=r["Timesheet Name"]
                new_match=r["Matched Rate Name"]
                new_conf=r["Confidence"]
                if IS_PG:
                    c.execute(
                        "UPDATE timesheet_entries SET matched_as=%s,ratio=%s WHERE name=%s",
                        (new_match,new_conf,name)
                    )
                else:
                    c.execute(
                        "UPDATE timesheet_entries SET matched_as=?,ratio=? WHERE name=?",
                        (new_match,new_conf,name)
                    )
            conn.commit()
            st.success(f"Updated {len(diffs)} match(es).")

# ---- 4) Dashboard ----
with tabs[3]:
    st.header("📊 Dashboard")
    st.markdown("…your existing charts here…")

# ---- 5) Settings ----
with tabs[4]:
    st.header("⚙️ Settings & Info")
    st.markdown("""
    - Duplicates by (name & date) are blocked.  
    - Bulk delete/export/mark‑paid in History.  
    - Inline edit matches in Matches.  
    - Custom or preset date‑ranges in History.  
    """)
